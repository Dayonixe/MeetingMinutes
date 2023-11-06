import os
import sys
import datetime
import tempfile
import math
import openai
from pydub import AudioSegment
from docx import Document

def split_audio(file_path):
    audio = AudioSegment.from_mp3(file_path)
    max_size = 25 * 1024 * 1024  # 25 MB
    chunk_length_ms = math.floor(0.9 * 1000 * (max_size / (audio.frame_rate * audio.frame_width)))
    chunks = [audio[i:i + chunk_length_ms] for i in range(0, len(audio), int(chunk_length_ms))]
    return chunks

def transcribe_audio(audio_chunks):
    transcriptions = []
    for chunk in audio_chunks:
        with tempfile.NamedTemporaryFile(suffix=".mp3", delete=True) as temp_audio_file:
            chunk.export(temp_audio_file.name, format="mp3", bitrate="192k")  # You can adjust the bitrate as needed
            file_size = os.path.getsize(temp_audio_file.name)
            if file_size > 25 * 1024 * 1024:
                raise ValueError("Audio chunk is too large: {} bytes".format(file_size))
            with open(temp_audio_file.name, 'rb') as f:
                transcription = openai.Audio.transcribe("whisper-1", f)
            transcriptions.append(transcription['text'])
    return " ".join(transcriptions)

def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Vous êtes une IA hautement qualifiée, formée à la compréhension et à la synthèse du langage. Sur la base du texte suivant, résumez en un paragraphe abstrait et concis. Essayez de retenir les points les plus importants, en fournissant un résumé cohérent et lisible qui pourrait aider une personne à comprendre les points principaux de la discussion sans avoir besoin de lire le texte en entier. Veuillez éviter les détails inutiles ou les points tangentiels."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Vous êtes une IA hautement qualifiée, spécialisé dans la distillation d'informations en points clés. Sur la base du texte suivant, identifiez et listez les points principaux qui ont été discutés ou évoqués. Il doit s'agir des idées, des résultats ou des sujets les plus importants qui sont cruciaux pour l'essence de la discussion. Votre objectif est de fournir une liste que quelqu'un pourrait lire pour comprendre rapidement ce qui a été discuté."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Vous êtes une IA hautement qualifiée, spécialisé dans l'analyse des conversations et de l'extraction des actions à entreprendre. Sur la base du texte suivant, identifiez les tâches, les missions ou les actions qui ont été convenues ou mentionnées comme devant être réalisées. Il peut s'agir de tâches assignées à des personnes spécifiques ou d'actions générales que le groupe a décidé d'entreprendre. Veuillez dresser une liste claire et concise de ces actions."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    return {
        'complete_transcription': transcription,
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items
    }

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python3 main.py <audio_file_path>")
        sys.exit(1)
    audio_file_path = sys.argv[1]
    openai.api_key = "API_KEY"

    # Ask the user if they want to run the full script or just the transcription
    while True:
        choice = input(
            "Voulez-vous exécuter le script complet (entrez 'F') ou seulement la transcription (entrez 'T') ? ").lower().strip()
        if choice in ['f', 't']:
            break
        else:
            print("Option non valide, veuillez saisir 'F' ou 'T'.")

    # Split audio into chunks
    audio_chunks = split_audio(audio_file_path)

    # Always perform transcription
    transcription = transcribe_audio(audio_chunks)

    # Check user's choice
    if choice == 'f':
        # If user chose 'full', perform all actions
        minutes = meeting_minutes(transcription)
        print("Minutes prepared.")
    elif choice == 't':
        # If user chose 'transcribe', only save the transcription
        minutes = {
            'complete_transcription': transcription
        }
        print("Transcription completed.")
    else:
        print("Invalid option. Exiting.")
        sys.exit(1)

    now = datetime.datetime.now()
    formatted_date = now.strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"meeting_minutes_{formatted_date}.docx"
    save_as_docx(minutes, filename)